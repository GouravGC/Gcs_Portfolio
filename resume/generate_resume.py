"""Generate the Master Resume (DOCX and optional PDF) for Gourav Chhatwani.

The resume is built from the shared structured data layer (data/profile.json and
data/projects.json) via resume.resume_template. No project content is duplicated
here — everything is assembled from the data layer.

Usage:
    python resume/generate_resume.py
"""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

# Ensure the project root is importable regardless of how the script is invoked.
_SCRIPT_DIR = Path(__file__).resolve().parent
_PROJECT_ROOT = _SCRIPT_DIR.parent
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

from utils.helpers import OUTPUT_DIR, ensure_output_dir

from resume.resume_template import build_resume_content

DOCX_NAME = "Gourav_Chhatwani_Master_Resume.docx"
PDF_NAME = "Gourav_Chhatwani_Master_Resume.pdf"


def _add_heading(doc, text: str, size: int = 24, color: str = "1F4E79"):
    """Add a heading paragraph with a bottom border."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Calibri"
    return p


def _add_section_title(doc, text: str):
    """Add a section title (e.g., TECHNICAL SKILLS)."""
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("1F4E79")
    run.font.name = "Calibri"
    return p


def _add_para(doc, text: str, bold: bool = False, size: int = 11, italic: bool = False):
    """Add a normal paragraph."""
    from docx.shared import Pt

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def _add_bullet(doc, text: str):
    """Add a bullet point."""
    from docx.shared import Pt

    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    return p


def generate_docx(content: dict) -> Path:
    """Generate the DOCX resume and return its path."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:  # pragma: no cover
        raise SystemExit(
            "python-docx is required. Install it with: pip install python-docx"
        ) from exc

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # ---------- Header ----------
    _add_heading(doc, content["name"], size=26)

    title_para = _add_para(doc, content["title"], bold=True, size=13)
    title_para.paragraph_format.space_after = Pt(4)

    # Contact line
    contact = " | ".join(
        [
            content.get("email", ""),
            content.get("phone", ""),
            content.get("location", ""),
            content.get("github", ""),
            content.get("linkedin", ""),
        ]
    )
    _add_para(doc, contact, size=10)

    # ---------- Summary ----------
    _add_section_title(doc, "Professional Summary")
    _add_para(doc, content["summary"], size=10.5)

    # ---------- Technical Skills ----------
    _add_section_title(doc, "Technical Skills")
    for line in content["skill_lines"]:
        _add_bullet(doc, line)

    # ---------- Projects ----------
    _add_section_title(doc, "Selected Projects")
    for entry in content["project_entries"]:
        _add_para(doc, entry["name"], bold=True, size=11)
        for bullet in entry["bullets"]:
            _add_bullet(doc, bullet)

    # ---------- Education ----------
    _add_section_title(doc, "Education")
    edu = content.get("education", [])
    if edu:
        for e in edu:
            _add_para(doc, e.get("degree", ""), bold=True, size=11)
            _add_para(doc, f"{e.get('institution', '')} — {e.get('year', '')}", size=10.5)
    else:
        _add_para(doc, "[Add education details]", italic=True, size=10.5)

    # ---------- Certifications ----------
    _add_section_title(doc, "Certifications")
    certs = content.get("certifications", [])
    if certs:
        for c in certs:
            _add_bullet(doc, f"{c.get('name', '')} — {c.get('issuer', '')}")
    else:
        _add_para(doc, "[Add certification details]", italic=True, size=10.5)

    # ---------- Links ----------
    _add_section_title(doc, "Links")
    _add_bullet(doc, f"GitHub: {content.get('github', '')}")
    _add_bullet(doc, f"LinkedIn: {content.get('linkedin', '')}")

    out_dir = ensure_output_dir()
    docx_path = out_dir / DOCX_NAME
    doc.save(str(docx_path))
    return docx_path


def convert_to_pdf(docx_path: Path) -> Path | None:
    """Attempt to convert DOCX to PDF using LibreOffice (soffice) if available."""
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUTPUT_DIR), str(docx_path)],
            check=True,
            capture_output=True,
            timeout=120,
        )
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return None
    return pdf_path if pdf_path.exists() else None


def main() -> None:
    """Build the resume content and generate DOCX / optional PDF."""
    print("Building resume content from data layer...")
    content = build_resume_content()

    print(f"Generating {DOCX_NAME}...")
    docx_path = generate_docx(content)
    print(f"  -> {docx_path}")

    pdf_path = convert_to_pdf(docx_path)
    if pdf_path:
        print(f"  -> {pdf_path}")
    else:
        print(
            "  PDF generation skipped (LibreOffice 'soffice' not available). "
            "The DOCX is ready."
        )

    print("Done.")


if __name__ == "__main__":
    main()
